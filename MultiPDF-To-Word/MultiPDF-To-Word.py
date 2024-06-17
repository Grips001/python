import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io
from tkinter import Tk, filedialog

def pdf_to_images(pdf_path):
    doc = fitz.open(pdf_path)
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)  # load page
        pix = page.get_pixmap()  # render page to an image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    return images

def add_images_to_doc(images, doc):
    for image in images:
        image_stream = io.BytesIO()
        image.save(image_stream, format='PNG')
        image_stream.seek(0)  # Ensure the stream position is at the beginning

        # Create a new paragraph for the image to control page breaks
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(image_stream, width=Inches(6))
        
        # Check if the image causes an automatic page break and adjust if necessary
        if len(doc.element.xpath('//w:lastRenderedPageBreak')) > 0:
            p._element.getparent().remove(p._element)
            doc.add_page_break()
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_picture(image_stream, width=Inches(6))
        doc.add_paragraph()  # Ensure separation of images

def add_toc(doc, toc_entries):
    toc = doc.add_paragraph()
    toc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = toc.add_run('Table of Contents')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    for toc_entry in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(toc_entry['title'])
        run.font.color.theme_color = 10  # Apply a color to indicate hyperlink
        r = run._r
        rPr = r.get_or_add_rPr()
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), toc_entry['anchor'])
        hyperlink.append(r)
        p._p.clear_content()
        p._p.append(hyperlink)

def merge_pdfs_to_word(pdf_files, output_word):
    doc = Document()
    toc_entries = []

    for idx, pdf_file in enumerate(pdf_files):
        title = f"Document {idx + 1}: {pdf_file.split('/')[-1]}"
        anchor = f"bookmark_{idx}"
        toc_entries.append({'title': title, 'anchor': anchor})

    # Add TOC at the beginning
    add_toc(doc, toc_entries)
    doc.add_page_break()

    for idx, pdf_file in enumerate(pdf_files):
        title = f"Document {idx + 1}: {pdf_file.split('/')[-1]}"
        anchor = f"bookmark_{idx}"

        # Create bookmarks and add images
        heading = doc.add_heading(title, level=1)

        # Add the bookmark start to the heading
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(idx))
        bookmark_start.set(qn('w:name'), anchor)
        heading._p.insert(0, bookmark_start)

        # Add the bookmark end after the images
        images = pdf_to_images(pdf_file)
        add_images_to_doc(images, doc)
        doc.add_page_break()

        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(idx))
        doc._body._element.append(bookmark_end)

    doc.save(output_word)

def select_pdf_files():
    root = Tk()
    root.withdraw()  # Hide the main window
    pdf_files = filedialog.askopenfilenames(
        title="Select PDF Files",
        filetypes=[("PDF Files", "*.pdf")]
    )
    return root.tk.splitlist(pdf_files)

def select_output_file():
    root = Tk()
    root.withdraw()  # Hide the main window
    output_file = filedialog.asksaveasfilename(
        title="Save Merged Document As",
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )
    return output_file

# Main script execution
if __name__ == "__main__":
    pdf_files = select_pdf_files()
    if not pdf_files:
        print("No PDF files selected. Exiting...")
    else:
        output_word = select_output_file()
        if not output_word:
            print("No output file selected. Exiting...")
        else:
            merge_pdfs_to_word(pdf_files, output_word)
            print(f"Merged document saved as {output_word}")
